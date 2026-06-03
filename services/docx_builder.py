# -*- coding: utf-8 -*-
"""DocxBuilder: Assemble the final Word document with proper formatting.

Directly reuses code from bid-tech-enrichment skill.
"""
import os
import sys
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

DEFAULT_TEMPLATE = r"C:\Users\shanzhao\Desktop\workbuddy投标书\商务+报价.docx"


def make_heading1(text):
    """Level 1 heading: 黑体 18pt bold"""
    p = etree.Element(qn('w:p'))
    pPr = etree.SubElement(p, qn('w:pPr'))
    pStyle = etree.SubElement(pPr, qn('w:pStyle'))
    pStyle.set(qn('w:val'), '1')
    r = etree.SubElement(p, qn('w:r'))
    rPr = etree.SubElement(r, qn('w:rPr'))
    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), '黑体')
    rFonts.set(qn('w:ascii'), 'SimHei')
    rFonts.set(qn('w:hAnsi'), 'SimHei')
    sz = etree.SubElement(rPr, qn('w:sz'))
    sz.set(qn('w:val'), '36')
    szCs = etree.SubElement(rPr, qn('w:szCs'))
    szCs.set(qn('w:val'), '36')
    b = etree.SubElement(rPr, qn('w:b'))
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    return p


def make_heading2(text):
    """Level 2 heading: 黑体 14pt bold"""
    p = etree.Element(qn('w:p'))
    pPr = etree.SubElement(p, qn('w:pPr'))
    pStyle = etree.SubElement(pPr, qn('w:pStyle'))
    pStyle.set(qn('w:val'), '2')
    r = etree.SubElement(p, qn('w:r'))
    rPr = etree.SubElement(r, qn('w:rPr'))
    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), '黑体')
    rFonts.set(qn('w:ascii'), 'SimHei')
    rFonts.set(qn('w:hAnsi'), 'SimHei')
    sz = etree.SubElement(rPr, qn('w:sz'))
    sz.set(qn('w:val'), '28')
    szCs = etree.SubElement(rPr, qn('w:szCs'))
    szCs.set(qn('w:val'), '28')
    b = etree.SubElement(rPr, qn('w:b'))
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    return p


def make_heading3(text):
    """Level 3 heading: 黑体 12pt bold"""
    p = etree.Element(qn('w:p'))
    pPr = etree.SubElement(p, qn('w:pPr'))
    pStyle = etree.SubElement(pPr, qn('w:pStyle'))
    pStyle.set(qn('w:val'), '3')
    r = etree.SubElement(p, qn('w:r'))
    rPr = etree.SubElement(r, qn('w:rPr'))
    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), '黑体')
    rFonts.set(qn('w:ascii'), 'SimHei')
    rFonts.set(qn('w:hAnsi'), 'SimHei')
    sz = etree.SubElement(rPr, qn('w:sz'))
    sz.set(qn('w:val'), '24')
    szCs = etree.SubElement(rPr, qn('w:szCs'))
    szCs.set(qn('w:val'), '24')
    b = etree.SubElement(rPr, qn('w:b'))
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    return p


def make_body(text):
    """Body text: 宋体 12pt, first-line indent 2 chars, fixed line spacing 20pt"""
    p = etree.Element(qn('w:p'))
    pPr = etree.SubElement(p, qn('w:pPr'))
    ind = etree.SubElement(pPr, qn('w:ind'))
    ind.set(qn('w:firstLine'), '480')
    spacing = etree.SubElement(pPr, qn('w:spacing'))
    spacing.set(qn('w:line'), '400')
    spacing.set(qn('w:lineRule'), 'exact')
    r = etree.SubElement(p, qn('w:r'))
    rPr = etree.SubElement(r, qn('w:rPr'))
    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'SimSun')
    rFonts.set(qn('w:hAnsi'), 'SimSun')
    sz = etree.SubElement(rPr, qn('w:sz'))
    sz.set(qn('w:val'), '24')
    szCs = etree.SubElement(rPr, qn('w:szCs'))
    szCs.set(qn('w:val'), '24')
    t = etree.SubElement(r, qn('w:t'))
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    return p


def parse_markdown_table(lines, start_idx):
    """Parse Markdown table lines, return (table_data, end_idx)."""
    table_data = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        if '---' in line.replace(' ', ''):
            i += 1
            continue
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            table_data.append(cells)
        i += 1
    return table_data, i


def make_table_element(headers, rows):
    """Create a Word table XML element with borders."""
    tbl = etree.Element(qn('w:tbl'))
    tblPr = etree.SubElement(tbl, qn('w:tblPr'))
    tblW = etree.SubElement(tblPr, qn('w:tblW'))
    tblW.set(qn('w:w'), '9000')
    tblW.set(qn('w:type'), 'dxa')
    tblLook = etree.SubElement(tblPr, qn('w:tblLook'))
    tblLook.set(qn('w:val'), '04A0')

    # Borders
    tblBorders = etree.SubElement(tblPr, qn('w:tblBorders'))
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = etree.SubElement(tblBorders, qn(f'w:{border_name}'))
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')

    tblGrid = etree.SubElement(tbl, qn('w:tblGrid'))
    col_width = 9000 // max(len(headers), 1)
    for _ in headers:
        gridCol = etree.SubElement(tblGrid, qn('w:gridCol'))
        gridCol.set(qn('w:w'), str(col_width))

    all_rows = [headers] + rows
    for row_data in all_rows:
        tr = etree.SubElement(tbl, qn('w:tr'))
        for cell_text in row_data:
            tc = etree.SubElement(tr, qn('w:tc'))
            tcPr = etree.SubElement(tc, qn('w:tcPr'))
            tcW = etree.SubElement(tcPr, qn('w:tcW'))
            tcW.set(qn('w:w'), str(col_width))
            tcW.set(qn('w:type'), 'dxa')
            # Cell margins
            tcMar = etree.SubElement(tcPr, qn('w:tcMar'))
            for m in ['top', 'left', 'bottom', 'right']:
                mar = etree.SubElement(tcMar, qn(f'w:{m}'))
                mar.set(qn('w:w'), '60')
                mar.set(qn('w:type'), 'dxa')
            p = etree.SubElement(tc, qn('w:p'))
            pPr = etree.SubElement(p, qn('w:pPr'))
            jc = etree.SubElement(pPr, qn('w:jc'))
            jc.set(qn('w:val'), 'center')
            r = etree.SubElement(p, qn('w:r'))
            rPr = etree.SubElement(r, qn('w:rPr'))
            rFonts = etree.SubElement(rPr, qn('w:rFonts'))
            rFonts.set(qn('w:eastAsia'), '宋体')
            rFonts.set(qn('w:ascii'), 'SimSun')
            sz = etree.SubElement(rPr, qn('w:sz'))
            sz.set(qn('w:val'), '24')
            t = etree.SubElement(r, qn('w:t'))
            t.text = cell_text
    return tbl


def insert_markdown_to_docx(doc, markdown_content: str):
    """Insert Markdown content into a python-docx Document object at {{TECH_SOLUTION_ANCHOR}}."""
    body = doc.element.body
    anchor_elem = None

    for child in body.iter():
        if child.tag == qn('w:t') and child.text and '{{TECH_SOLUTION_ANCHOR}}' in child.text:
            p_elem = child.getparent().getparent()
            if p_elem.tag == qn('w:p'):
                anchor_elem = p_elem
                break

    if anchor_elem is None:
        raise ValueError("Template missing {{TECH_SOLUTION_ANCHOR}} placeholder")

    lines = markdown_content.strip().split('\n')
    new_elements = []
    i = 0

    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith('|'):
            table_data, end_i = parse_markdown_table(lines, i)
            if len(table_data) >= 2:
                headers = table_data[0]
                rows = table_data[1:]
                tbl_elem = make_table_element(headers, rows)
                new_elements.append(tbl_elem)
                i = end_i
                continue
            else:
                new_elements.append(make_body(stripped))
                i += 1
                continue

        if stripped.startswith('# ') and not stripped.startswith('## '):
            new_elements.append(make_heading1(stripped[2:]))
        elif stripped.startswith('## ') and not stripped.startswith('### '):
            new_elements.append(make_heading2(stripped[3:]))
        elif stripped.startswith('### '):
            new_elements.append(make_heading3(stripped[4:]))
        else:
            new_elements.append(make_body(stripped))

        i += 1

    prev = anchor_elem
    for elem in new_elements:
        prev.addnext(elem)
        prev = elem

    body.remove(anchor_elem)


def unify_all_paragraphs(doc):
    """Apply formatting to all paragraphs based on heading style."""
    body = doc.element.body

    for para_elem in body.iter(qn('w:p')):
        pPr = para_elem.find(qn('w:pPr'))
        pStyle = pPr.find(qn('w:pStyle')) if pPr is not None else None
        style_val = pStyle.get(qn('w:val')) if pStyle is not None else None

        if style_val in ('1', 'Heading1', 'heading 1'):
            _format_heading_para(para_elem, 18, 0)
        elif style_val in ('2', 'Heading2', 'heading 2'):
            _format_heading_para(para_elem, 14, 1)
        elif style_val in ('3', 'Heading3', 'heading 3'):
            _format_heading_para(para_elem, 12, 2)
        else:
            _format_body_para(para_elem)


def _format_heading_para(para_elem, font_size_pt, outline_level):
    """Format heading: 黑体 + size + bold + outline level."""
    pPr = para_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = etree.Element(qn('w:pPr'))
        para_elem.insert(0, pPr)

    for ind in pPr.findall(qn('w:ind')):
        pPr.remove(ind)
    for sp in pPr.findall(qn('w:spacing')):
        pPr.remove(sp)

    outline = pPr.find(qn('w:outlineLvl'))
    if outline is None:
        outline = etree.SubElement(pPr, qn('w:outlineLvl'))
    outline.set(qn('w:val'), str(outline_level))

    half_pts = str(font_size_pt * 2)
    for r_elem in para_elem.iter(qn('w:r')):
        rPr = r_elem.find(qn('w:rPr'))
        if rPr is None:
            rPr = etree.Element(qn('w:rPr'))
            r_elem.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        rFonts.set(qn('w:eastAsia'), '黑体')
        rFonts.set(qn('w:ascii'), 'SimHei')
        rFonts.set(qn('w:hAnsi'), 'SimHei')

        sz = rPr.find(qn('w:sz'))
        if sz is None:
            sz = etree.SubElement(rPr, qn('w:sz'))
        sz.set(qn('w:val'), half_pts)
        szCs = rPr.find(qn('w:szCs'))
        if szCs is None:
            szCs = etree.SubElement(rPr, qn('w:szCs'))
        szCs.set(qn('w:val'), half_pts)

        b = rPr.find(qn('w:b'))
        if b is None:
            etree.SubElement(rPr, qn('w:b'))


def _format_body_para(para_elem):
    """Format body: 宋体 12pt, indent, spacing."""
    pPr = para_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = etree.Element(qn('w:pPr'))
        para_elem.insert(0, pPr)

    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = etree.SubElement(pPr, qn('w:ind'))
    ind.set(qn('w:firstLine'), '480')

    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = etree.SubElement(pPr, qn('w:spacing'))
    spacing.set(qn('w:line'), '400')
    spacing.set(qn('w:lineRule'), 'exact')

    outline = pPr.find(qn('w:outlineLvl'))
    if outline is not None:
        pPr.remove(outline)

    for r_elem in para_elem.iter(qn('w:r')):
        rPr = r_elem.find(qn('w:rPr'))
        if rPr is None:
            rPr = etree.Element(qn('w:rPr'))
            r_elem.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        rFonts.set(qn('w:eastAsia'), '宋体')
        rFonts.set(qn('w:ascii'), 'SimSun')
        rFonts.set(qn('w:hAnsi'), 'SimSun')

        sz = rPr.find(qn('w:sz'))
        if sz is None:
            sz = etree.SubElement(rPr, qn('w:sz'))
        sz.set(qn('w:val'), '24')
        szCs = rPr.find(qn('w:szCs'))
        if szCs is None:
            szCs = etree.SubElement(rPr, qn('w:szCs'))
        szCs.set(qn('w:val'), '24')

        b = rPr.find(qn('w:b'))
        if b is not None:
            rPr.remove(b)


def replace_first_page_title(doc, project_name):
    """Replace first page default titles with project name."""
    body = doc.element.body
    para_count = 0

    for para_elem in body.iter(qn('w:p')):
        para_count += 1
        if para_count > 30:
            break

        texts = []
        run_elements = []
        for r_elem in para_elem.iter(qn('w:r')):
            for t_elem in r_elem.iter(qn('w:t')):
                if t_elem.text:
                    texts.append(t_elem.text)
                    run_elements.append(t_elem)

        full_text = ''.join(texts)
        has_keyword = any(k in full_text for k in ['商务文件', '报价文件', '商务部分', '报价部分', '商务', '报价'])

        pPr = para_elem.find(qn('w:pPr'))
        is_title = False
        if pPr is not None:
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                style_val = pStyle.get(qn('w:val'), '')
                is_title = '1' in style_val or 'Title' in style_val or 'Heading' in style_val

        if (has_keyword or is_title) and full_text.strip():
            for run_elem in run_elements:
                run_elem.text = ''
            if run_elements:
                run_elements[0].text = project_name
                r_elem = run_elements[0].getparent()
                rPr = r_elem.find(qn('w:rPr'))
                if rPr is None:
                    rPr = etree.Element(qn('w:rPr'))
                    r_elem.insert(0, rPr)
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = etree.SubElement(rPr, qn('w:rFonts'))
                rFonts.set(qn('w:eastAsia'), '黑体')
                rFonts.set(qn('w:ascii'), 'SimHei')
                sz = rPr.find(qn('w:sz'))
                if sz is None:
                    sz = etree.SubElement(rPr, qn('w:sz'))
                sz.set(qn('w:val'), '44')
                b = rPr.find(qn('w:b'))
                if b is None:
                    etree.SubElement(rPr, qn('w:b'))
                if pPr is not None:
                    jc = pPr.find(qn('w:jc'))
                    if jc is None:
                        jc = etree.SubElement(pPr, qn('w:jc'))
                    jc.set(qn('w:val'), 'center')
            break


def remove_empty_paragraphs(doc):
    """Remove blank paragraphs that are direct children of body (skip table cells)."""
    body = doc.element.body
    to_remove = []

    # Only iterate direct children of body to avoid touching table-cell paragraphs
    for child in body:
        if child.tag == qn('w:p'):
            all_text = ''
            for t_elem in child.iter(qn('w:t')):
                if t_elem.text:
                    all_text += t_elem.text
            if not all_text.strip() or '{{TECH_SOLUTION_ANCHOR}}' in all_text:
                to_remove.append(child)

    for para_elem in to_remove:
        body.remove(para_elem)


def inject_anchor_if_needed(template_path: str) -> str:
    """Ensure template has {{TECH_SOLUTION_ANCHOR}} placeholder. Returns path to use."""
    doc = Document(template_path)
    full_text = "\n".join([p.text for p in doc.paragraphs])

    if "{{TECH_SOLUTION_ANCHOR}}" in full_text:
        return template_path

    # Find insertion point before 第三卷
    insert_index = None
    for i, para in enumerate(doc.paragraphs):
        if "第三卷" in para.text:
            insert_index = i
            break

    if insert_index is not None:
        anchor_para = doc.paragraphs[insert_index]._element
        new_p = doc.paragraphs[insert_index]._element.makeelement(qn('w:p'), {})
        r_elem = new_p.makeelement(qn('w:r'), {})
        t_elem = r_elem.makeelement(qn('w:t'), {})
        t_elem.text = "{{TECH_SOLUTION_ANCHOR}}"
        r_elem.append(t_elem)
        new_p.append(r_elem)
        anchor_para.addprevious(new_p)

    output_copy = template_path.replace(".docx", "_带占位符.docx")
    doc.save(output_copy)
    return output_copy


def build_final_docx(
    expanded_solution: str,
    project_name: str,
    output_path: str,
    template_path: str = None,
):
    """Complete assembly pipeline: inject anchor, insert tech, format, save."""
    template_path = template_path or DEFAULT_TEMPLATE

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    # Step 1: Ensure anchor
    working_template = inject_anchor_if_needed(template_path)

    # Step 2: Load and insert
    doc = Document(working_template)
    insert_markdown_to_docx(doc, expanded_solution)

    # Step 3: Format
    unify_all_paragraphs(doc)
    replace_first_page_title(doc, project_name)
    remove_empty_paragraphs(doc)

    # Step 4: Save
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    # Clean up temp file
    if working_template != template_path and os.path.exists(working_template):
        try:
            os.remove(working_template)
        except Exception:
            pass

    return output_path
